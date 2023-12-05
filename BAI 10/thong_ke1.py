# ---------------------------------
# Code ban đầu
# ---------------------------------

# import pandas as pd
# from numpy import array
# import matplotlib.pyplot as plt
# import numpy as np
# import csv

# df=pd.read_csv('diemPython.csv',index_col=0,header = 0)
# in_data = array(df.iloc[:,:])
# print(in_data)
# print('Tong so sinh vien di thi :')
# tongsv= in_data[:,1]
# print(np.sum(tongsv))
# diemA = in_data[:,3]
# diemBc = in_data[:,4]
# print('Tong sv:',tongsv)
# maxa = diemA.max()
# i, = np.where(diemA == maxa)
# print('lop co nhieu diem A la {0} co {1} sv dat diem A'.format(in_data[i,0],maxa))
# plt.plot(range(len(diemA)),diemA,'r-',label="Diem A")
# plt.plot(range(len(diemBc)),diemBc,'g-',label="Diem B +")
# plt.xlabel('Lơp')
# plt.ylabel(' So sv dat diem ')
# plt.legend(loc='upper right')
# plt.show()
      
# ---------------------------------
# Copyright by Quynh Trang Teacher
# ---------------------------------

from tkinter import *
from tkinter import ttk
from PIL import ImageTk, Image
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm
from openpyxl.utils import get_column_letter



# ---------------------------------
# Giao diện
# ---------------------------------
class App:
    def __init__(self):
        self.root = Tk()
        self.root.title('Final Report')
        self.root.geometry('1000x700')

        # header text style
        self.header_font = ("Times New Roman", 20, 'bold')
        self.text_font = ("Times New Roman", 14, 'normal')

        # biểu đồ, tỉ lệ ảnh là 4:3
        self.score_ratio = ImageTk.PhotoImage(Image.open('score_ratio_1.png').resize((400,300)))
        self.test_ratio = ImageTk.PhotoImage(Image.open('test_ratio.png').resize((400,300)))

        self.style = ttk.Style()
        self.style.configure("TFrame", background="#fff")
        
        self.style1 = ttk.Style()
        self.style1.configure("Frame2.TFrame",background="#f0f0f0")

    

    # ---------------------------------
    # Copyright by Nguyen Van Long
    # ---------------------------------
    def export_word(self):
        # Tạo một tài liệu mới
        doc = Document()

        # Đặt kích thước giấy A4
        section = doc.sections[0]
        section.page_width = Cm(21) # Chiều rộng A4 (21 cm)
        section.page_height = Cm(29.7) # height 29.7 cm

        # Định dạng tiêu đề
        title = doc.add_paragraph("Báo cáo môn học Lập trình Python (FE6051)")
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(18)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        body_text = [
            "I.Thông tin chung","Tên môn: Lập trình Python",
            "Mã môn: FE6051        Số tín chỉ: 3(2,1,0)",
            "Số lớp học phần: 9",
            "Tổng số sinh viên: 519        Pass: 83.2%",
            "II. Kết quả xử lý số liệu"
        ]
        
        for i in body_text:
            body_para = doc.add_paragraph(i)
            body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            body_para.paragraph_format.line_spacing = Pt(15) # dãn dòng 1.5
            body_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            body_para.runs[0].font.size = Pt(14)

        # Chèn hình ảnh
        doc.add_picture('./score_ratio_1.png')
        doc.add_picture('./test_ratio.png')

        end_text = [
            "III. Kết luận",
            "."*79,
            "."*79
        ]

        for i in end_text:
            end_para = doc.add_paragraph(i)
            end_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            end_para.paragraph_format.line_spacing = Pt(15) # dãn dòng 1.5
            end_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            end_para.runs[0].font.size = Pt(14)

        footer = "Ngày .... Tháng .... Năm .... \n\nKý tên\t\t"
        footer_doc = doc.add_paragraph(footer)
        footer_doc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_doc.paragraph_format.line_spacing = Pt(15) # dãn dòng 1.5
        footer_doc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        footer_doc.runs[0].font.size = Pt(14)

        # Đặt font family cho toàn tài liệu
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"


        # Lưu tài liệu
        doc.save("Final_report.docx")
    # ---------------------------------
    # Copyright by Nguyen Van Long
    # ---------------------------------





    # làm cho GUI khi hiển thị sẽ căn giữa màn hình
    def make_center(self):
        self.root.update_idletasks()
        # lấy chiều rộng và cao của window tkinter
        width = self.root.winfo_width()
        height = self.root.winfo_height()

        # lấy chiều rộng màn hình laptop chia 2 trừ đi chiều rộng cửa sổ tkinter / 2
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry("{}x{}+{}+{}".format(width,height,x,y))


# ---------------------------------
# Copyright by Nguyen Van Long & Tran Ba Quang
# ---------------------------------

if __name__ == '__main__':
    app = App()
    app.make_center()
    app.draw_ui()
    app.root.mainloop()